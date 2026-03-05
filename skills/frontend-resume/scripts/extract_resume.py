#!/usr/bin/env python3
"""
Resume Text Extractor for frontend-resume skill.

Usage:
    python extract_resume.py <resume_file> [--format json|text]

Supported formats: .pdf, .docx, .md, .txt, .json (JSON Resume)

Extracts text with structural hints (font size, bold) for downstream
resume section parsing. Output is a list of pages, each containing lines
with text, font_size, and is_bold metadata.
"""

import sys
import json
import os
import re


# ---------------------------------------------------------------------------
# PDF extractors (pdfplumber primary, PyPDF2 fallback)
# ---------------------------------------------------------------------------

def extract_with_pdfplumber(file_path):
    """Extract text using pdfplumber with layout and font size hints."""
    import pdfplumber

    pages_data = []
    with pdfplumber.open(file_path) as pdf:
        for page_num, page in enumerate(pdf.pages):
            lines = []
            chars = page.chars
            if chars:
                current_line = []
                current_top = None
                for char in chars:
                    top = round(char["top"], 1)
                    if current_top is None:
                        current_top = top
                    if abs(top - current_top) > 3:
                        if current_line:
                            text = "".join(c["text"] for c in current_line).strip()
                            if text:
                                avg_size = sum(c.get("size", 12) for c in current_line) / len(current_line)
                                is_bold = any(
                                    "Bold" in str(c.get("fontname", "")) or "bold" in str(c.get("fontname", ""))
                                    for c in current_line
                                )
                                lines.append({
                                    "text": text,
                                    "font_size": round(avg_size, 1),
                                    "is_bold": is_bold,
                                })
                        current_line = [char]
                        current_top = top
                    else:
                        current_line.append(char)
                if current_line:
                    text = "".join(c["text"] for c in current_line).strip()
                    if text:
                        avg_size = sum(c.get("size", 12) for c in current_line) / len(current_line)
                        is_bold = any(
                            "Bold" in str(c.get("fontname", "")) or "bold" in str(c.get("fontname", ""))
                            for c in current_line
                        )
                        lines.append({
                            "text": text,
                            "font_size": round(avg_size, 1),
                            "is_bold": is_bold,
                        })
            else:
                raw_text = page.extract_text()
                if raw_text:
                    for line in raw_text.split("\n"):
                        stripped = line.strip()
                        if stripped:
                            lines.append({
                                "text": stripped,
                                "font_size": 12.0,
                                "is_bold": False,
                            })

            pages_data.append({
                "page": page_num + 1,
                "lines": lines,
            })

    return pages_data


def extract_with_pypdf2(file_path):
    """Extract text using PyPDF2 as fallback (no font info)."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages_data = []

    if reader.is_encrypted:
        try:
            reader.decrypt("")
        except Exception:
            print("ERROR: PDF is password-protected. Please provide an unencrypted PDF.", file=sys.stderr)
            sys.exit(1)

    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        lines = []
        for line in text.split("\n"):
            stripped = line.strip()
            if stripped:
                lines.append({
                    "text": stripped,
                    "font_size": 12.0,
                    "is_bold": False,
                })
        pages_data.append({
            "page": page_num + 1,
            "lines": lines,
        })

    return pages_data


def extract_pdf(file_path):
    """Extract text from PDF, trying pdfplumber then PyPDF2."""
    pages_data = None
    extractor_used = None

    try:
        pages_data = extract_with_pdfplumber(file_path)
        extractor_used = "pdfplumber"
    except ImportError:
        try:
            pages_data = extract_with_pypdf2(file_path)
            extractor_used = "PyPDF2"
        except ImportError:
            print("ERROR: No PDF library available.", file=sys.stderr)
            print("Install one of:", file=sys.stderr)
            print("  pip install pdfplumber", file=sys.stderr)
            print("  pip install PyPDF2", file=sys.stderr)
            sys.exit(1)
    except Exception as e:
        print(f"ERROR: pdfplumber failed: {e}", file=sys.stderr)
        try:
            pages_data = extract_with_pypdf2(file_path)
            extractor_used = "PyPDF2 (fallback)"
        except ImportError:
            print("ERROR: PyPDF2 not available as fallback.", file=sys.stderr)
            print("  pip install PyPDF2", file=sys.stderr)
            sys.exit(1)

    total_text = sum(len(line["text"]) for page in pages_data for line in page["lines"])
    if total_text < 50:
        print("WARNING: Very little text extracted. The PDF may be image-based (scanned).", file=sys.stderr)
        print("Consider using an OCR tool first, or provide resume content as text.", file=sys.stderr)

    print(f"# Extracted with: {extractor_used}", file=sys.stderr)
    return pages_data


# ---------------------------------------------------------------------------
# DOCX extractor
# ---------------------------------------------------------------------------

HEADING_FONT_SIZES = {
    "Heading 1": 24.0,
    "Heading 2": 18.0,
    "Heading 3": 14.0,
    "Heading 4": 13.0,
    "Title": 26.0,
    "Subtitle": 18.0,
}


def extract_docx(file_path):
    """Extract text from DOCX with paragraph style and bold hints."""
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx is required for .docx files.", file=sys.stderr)
        print("  pip install python-docx", file=sys.stderr)
        sys.exit(1)

    doc = Document(file_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else "Normal"
        font_size = HEADING_FONT_SIZES.get(style_name, 12.0)

        is_bold = False
        for run in para.runs:
            if run.bold:
                is_bold = True
                break
            if run.font and run.font.size:
                pt = run.font.size.pt
                if pt > font_size:
                    font_size = pt

        lines.append({
            "text": text,
            "font_size": font_size,
            "is_bold": is_bold or style_name.startswith("Heading"),
        })

    for table in doc.tables:
        for row in table.rows:
            cells_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells_text:
                lines.append({
                    "text": " | ".join(cells_text),
                    "font_size": 12.0,
                    "is_bold": False,
                })

    print("# Extracted with: python-docx", file=sys.stderr)
    return [{"page": 1, "lines": lines}]


# ---------------------------------------------------------------------------
# Markdown extractor
# ---------------------------------------------------------------------------

def _strip_md_inline(text):
    """Remove inline Markdown formatting, convert links to 'text (url)'."""
    text = re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    text = re.sub(r'`(.+?)`', r'\1', text)
    return text.strip()


def extract_markdown(file_path):
    """Extract text from Markdown with heading-level hints."""
    with open(file_path, "r", encoding="utf-8") as f:
        raw = f.read()

    lines = []
    for raw_line in raw.split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            continue

        # Skip YAML frontmatter delimiter
        if stripped == "---":
            continue

        heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            size_map = {1: 24.0, 2: 18.0, 3: 14.0, 4: 13.0, 5: 12.0, 6: 12.0}
            lines.append({
                "text": _strip_md_inline(heading_match.group(2)),
                "font_size": size_map.get(level, 12.0),
                "is_bold": True,
            })
            continue

        has_bold = bool(re.search(r'\*\*(.+?)\*\*|__(.+?)__', stripped))
        clean = _strip_md_inline(stripped)

        list_match = re.match(r'^[-*+]\s+(.*)$', clean)
        if list_match:
            clean = "• " + list_match.group(1)

        num_match = re.match(r'^\d+\.\s+(.*)$', clean)
        if num_match:
            clean = re.sub(r'^(\d+\.)\s+', r'\1 ', clean)

        if clean:
            lines.append({
                "text": clean,
                "font_size": 12.0,
                "is_bold": has_bold,
            })

    print("# Extracted with: markdown-parser", file=sys.stderr)
    return [{"page": 1, "lines": lines}]


# ---------------------------------------------------------------------------
# Plain text extractor
# ---------------------------------------------------------------------------

def extract_plaintext(file_path):
    """Extract text from plain text. ALL-CAPS lines are treated as headings."""
    with open(file_path, "r", encoding="utf-8") as f:
        raw = f.read()

    lines = []
    for raw_line in raw.split("\n"):
        stripped = raw_line.strip()
        if not stripped:
            continue

        alpha_chars = re.sub(r'[^a-zA-Z]', '', stripped)
        is_allcaps = len(alpha_chars) >= 3 and alpha_chars == alpha_chars.upper()

        lines.append({
            "text": stripped,
            "font_size": 16.0 if is_allcaps else 12.0,
            "is_bold": is_allcaps,
        })

    print("# Extracted with: plaintext-reader", file=sys.stderr)
    return [{"page": 1, "lines": lines}]


# ---------------------------------------------------------------------------
# JSON Resume extractor (jsonresume.org schema)
# ---------------------------------------------------------------------------

def _jr_line(text, font_size=12.0, is_bold=False):
    return {"text": text, "font_size": font_size, "is_bold": is_bold}


def extract_json_resume(file_path):
    """Extract text from a JSON Resume (jsonresume.org schema)."""
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)

    lines = []

    basics = data.get("basics", {})
    if basics.get("name"):
        lines.append(_jr_line(basics["name"], 24.0, True))
    if basics.get("label"):
        lines.append(_jr_line(basics["label"], 14.0))
    contact_parts = []
    if basics.get("email"):
        contact_parts.append(basics["email"])
    if basics.get("phone"):
        contact_parts.append(basics["phone"])
    if basics.get("url"):
        contact_parts.append(basics["url"])
    location = basics.get("location", {})
    if isinstance(location, dict):
        loc_parts = [location.get("city", ""), location.get("region", ""), location.get("countryCode", "")]
        loc_str = ", ".join(p for p in loc_parts if p)
        if loc_str:
            contact_parts.append(loc_str)
    if contact_parts:
        lines.append(_jr_line(" | ".join(contact_parts), 12.0))
    for profile in basics.get("profiles", []):
        network = profile.get("network", "")
        username = profile.get("username", "")
        url = profile.get("url", "")
        lines.append(_jr_line(f"{network}: {username} ({url})" if url else f"{network}: {username}", 12.0))
    if basics.get("summary"):
        lines.append(_jr_line("Summary", 18.0, True))
        lines.append(_jr_line(basics["summary"], 12.0))

    for work in data.get("work", []):
        if not lines or lines[-1].get("font_size") != 18.0 or lines[-1].get("text") != "Experience":
            if data.get("work") and data["work"][0] is work:
                lines.append(_jr_line("Experience", 18.0, True))
        role = work.get("position", work.get("name", ""))
        company = work.get("name", work.get("company", ""))
        dates = f"{work.get('startDate', '')} - {work.get('endDate', 'Present')}"
        lines.append(_jr_line(f"{role} @ {company}", 14.0, True))
        lines.append(_jr_line(dates, 12.0))
        if work.get("summary"):
            lines.append(_jr_line(work["summary"], 12.0))
        for hl in work.get("highlights", []):
            lines.append(_jr_line(f"• {hl}", 12.0))

    if data.get("education"):
        lines.append(_jr_line("Education", 18.0, True))
        for edu in data["education"]:
            institution = edu.get("institution", "")
            area = edu.get("area", "")
            study_type = edu.get("studyType", "")
            degree = f"{study_type} in {area}" if study_type and area else (study_type or area)
            dates = f"{edu.get('startDate', '')} - {edu.get('endDate', '')}"
            lines.append(_jr_line(f"{degree} — {institution}", 14.0, True))
            lines.append(_jr_line(dates, 12.0))
            if edu.get("score"):
                lines.append(_jr_line(f"GPA: {edu['score']}", 12.0))
            for course in edu.get("courses", []):
                lines.append(_jr_line(f"• {course}", 12.0))

    if data.get("skills"):
        lines.append(_jr_line("Skills", 18.0, True))
        for skill in data["skills"]:
            name = skill.get("name", "")
            keywords = ", ".join(skill.get("keywords", []))
            level = skill.get("level", "")
            parts = [name]
            if level:
                parts.append(f"({level})")
            if keywords:
                parts.append(f": {keywords}")
            lines.append(_jr_line(" ".join(parts), 12.0))

    if data.get("projects"):
        lines.append(_jr_line("Projects", 18.0, True))
        for proj in data["projects"]:
            lines.append(_jr_line(proj.get("name", ""), 14.0, True))
            if proj.get("description"):
                lines.append(_jr_line(proj["description"], 12.0))
            if proj.get("keywords"):
                lines.append(_jr_line("Tech: " + ", ".join(proj["keywords"]), 12.0))
            if proj.get("url"):
                lines.append(_jr_line(proj["url"], 12.0))
            for hl in proj.get("highlights", []):
                lines.append(_jr_line(f"• {hl}", 12.0))

    if data.get("certificates"):
        lines.append(_jr_line("Certifications", 18.0, True))
        for cert in data["certificates"]:
            name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            lines.append(_jr_line(f"{name} — {issuer} ({date})" if date else f"{name} — {issuer}", 12.0))

    if data.get("languages"):
        lines.append(_jr_line("Languages", 18.0, True))
        for lang in data["languages"]:
            name = lang.get("language", "")
            fluency = lang.get("fluency", "")
            lines.append(_jr_line(f"{name} ({fluency})" if fluency else name, 12.0))

    if data.get("awards"):
        lines.append(_jr_line("Awards", 18.0, True))
        for award in data["awards"]:
            title = award.get("title", "")
            awarder = award.get("awarder", "")
            date = award.get("date", "")
            lines.append(_jr_line(f"{title} — {awarder} ({date})" if date else f"{title} — {awarder}", 12.0, True))
            if award.get("summary"):
                lines.append(_jr_line(award["summary"], 12.0))

    if data.get("volunteer"):
        lines.append(_jr_line("Volunteer", 18.0, True))
        for vol in data["volunteer"]:
            org = vol.get("organization", "")
            position = vol.get("position", "")
            dates = f"{vol.get('startDate', '')} - {vol.get('endDate', 'Present')}"
            lines.append(_jr_line(f"{position} @ {org}", 14.0, True))
            lines.append(_jr_line(dates, 12.0))
            if vol.get("summary"):
                lines.append(_jr_line(vol["summary"], 12.0))
            for hl in vol.get("highlights", []):
                lines.append(_jr_line(f"• {hl}", 12.0))

    if data.get("publications"):
        lines.append(_jr_line("Publications", 18.0, True))
        for pub in data["publications"]:
            name = pub.get("name", "")
            publisher = pub.get("publisher", "")
            date = pub.get("releaseDate", "")
            lines.append(_jr_line(f"{name} — {publisher} ({date})" if date else f"{name} — {publisher}", 14.0, True))
            if pub.get("summary"):
                lines.append(_jr_line(pub["summary"], 12.0))
            if pub.get("url"):
                lines.append(_jr_line(pub["url"], 12.0))

    print("# Extracted with: json-resume-parser", file=sys.stderr)
    return [{"page": 1, "lines": lines}]


# ---------------------------------------------------------------------------
# Output formatting
# ---------------------------------------------------------------------------

def format_as_text(pages_data):
    """Format extracted data as plain text with structural hints."""
    output = []
    for page in pages_data:
        output.append(f"--- Page {page['page']} ---")
        for line in page["lines"]:
            prefix = ""
            if line["font_size"] > 14:
                prefix = "[H] "
            elif line["is_bold"]:
                prefix = "[B] "
            output.append(f"{prefix}{line['text']}")
        output.append("")
    return "\n".join(output)


# ---------------------------------------------------------------------------
# Main dispatcher
# ---------------------------------------------------------------------------

SUPPORTED_EXTENSIONS = {
    ".pdf": "PDF",
    ".docx": "Word (DOCX)",
    ".md": "Markdown",
    ".txt": "Plain Text",
    ".json": "JSON Resume",
}


def main():
    if len(sys.argv) < 2:
        supported = ", ".join(SUPPORTED_EXTENSIONS.keys())
        print(f"Usage: python extract_resume.py <resume_file> [--format json|text]", file=sys.stderr)
        print(f"Supported formats: {supported}", file=sys.stderr)
        sys.exit(1)

    file_path = sys.argv[1]
    output_format = "text"

    if "--format" in sys.argv:
        idx = sys.argv.index("--format")
        if idx + 1 < len(sys.argv):
            output_format = sys.argv[idx + 1]

    if not os.path.exists(file_path):
        print(f"ERROR: File not found: {file_path}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(file_path)[1].lower()

    extractors = {
        ".pdf": extract_pdf,
        ".docx": extract_docx,
        ".md": extract_markdown,
        ".txt": extract_plaintext,
        ".json": extract_json_resume,
    }

    extractor = extractors.get(ext)
    if not extractor:
        supported = ", ".join(SUPPORTED_EXTENSIONS.keys())
        print(f"ERROR: Unsupported file format '{ext}'.", file=sys.stderr)
        print(f"Supported formats: {supported}", file=sys.stderr)
        sys.exit(1)

    pages_data = extractor(file_path)

    total_text = sum(len(line["text"]) for page in pages_data for line in page["lines"])
    print(f"# Pages: {len(pages_data)}, Total chars: {total_text}", file=sys.stderr)

    if output_format == "json":
        print(json.dumps(pages_data, ensure_ascii=False, indent=2))
    else:
        print(format_as_text(pages_data))


if __name__ == "__main__":
    main()
